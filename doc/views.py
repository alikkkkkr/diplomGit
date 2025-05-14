import os
from io import BytesIO
from urllib.parse import quote

from django.core.mail import EmailMultiAlternatives
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.template.loader import render_to_string
from docx.shared import Pt
from .forms import *
import pandas as pd
from .models import *
from django.http import JsonResponse, Http404, FileResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
import json
from django.core.paginator import Paginator
from docx import Document as DocxDocument
import time


def upload_interns(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        try:
            # Проверка авторизации пользователя
            if not request.session.get('email'):
                messages.error(request, "Для загрузки файла необходимо авторизоваться")
                return redirect('authPage')

            # Получаем текущего пользователя
            user = Account.objects.get(email=request.session['email'])

            # Чтение файла Excel
            excel_file = request.FILES['excel_file']
            df = pd.read_excel(excel_file, sheet_name=None)  # Читаем все листы

            for sheet_name, sheet_data in df.items():
                # Извлекаем название группы из имени листа
                group_name_full = sheet_name.strip()  # Имя листа
                group_name = group_name_full.replace("Группа ", "").strip()  # Убираем слово "Группа ", если есть

                # Проверяем, существует ли уже такая группа
                existing_group = Group.objects.filter(name=group_name).first()
                if existing_group:
                    # Если группа уже существует, проверяем, привязана ли она к пользователю
                    if user.role.name == 'Руководитель практики' and existing_group not in user.managed_groups.all():
                        user.managed_groups.add(existing_group)
                        user.save()
                        messages.info(request, f"Группа {group_name} уже существует и была добавлена к вашему профилю")
                    continue  # Пропускаем обработку существующей группы

                # Определяем код специальности на основе названия группы
                specialty_code = None
                if group_name.startswith(("П50", "П")):  # Программисты
                    specialty_code = "09.02.07"
                elif group_name.startswith(("Э50", "Э")):  # Программисты
                    specialty_code = "09.02.01"
                else:
                    raise ValueError(f"Неизвестная группа: {group_name}. Не удалось определить специальность.")

                # Находим специальность по коду
                specialty = Specialty.objects.filter(code=specialty_code).first()
                if not specialty:
                    raise ValueError(f"Специальность с кодом '{specialty_code}' не найдена в базе данных.")

                # Создаем новую группу с привязкой к специальности
                group = Group.objects.create(name=group_name, specialty=specialty)

                # Если пользователь - руководитель практики, добавляем группу к его управляемым группам
                if user.role.name == 'Руководитель практики':
                    user.managed_groups.add(group)
                    user.save()
                    messages.success(request, f"Группа {group_name} создана и добавлена к вашему профилю")

                # Явно указываем индексы столбцов
                column_mapping = {
                    "Фамилия Имя Отчество": 1,  # Столбец V (индекс 1)
                    "Ближайшее метро от дома": 2,  # Столбец C (индекс 2)
                    "Мобильный телефон": 3,  # Столбец D (индекс 3)
                    "База практики": 4,  # Столбец E (индекс 4)
                }

                # Проверяем, что данные существуют в соответствующих столбцах
                for col_name, col_idx in column_mapping.items():
                    if col_idx >= len(sheet_data.columns):
                        raise ValueError(f"Отсутствует столбец для поля: {col_name}")

                # Начинаем обработку данных с третьей строки (индекс 2)
                data_rows = sheet_data.iloc[2:]  # Начинаем с строки V3

                # Обработка строк таблицы
                for _, row in data_rows.iterrows():
                    # ФИО практиканта
                    full_name = row.iloc[column_mapping["Фамилия Имя Отчество"]]
                    if not full_name or pd.isna(full_name):  # Проверяем на пустоту или NaN
                        continue  # Пропускаем строку, если ФИО отсутствует

                    full_name_parts = full_name.split()
                    if len(full_name_parts) < 2:
                        continue  # Пропускаем строку, если ФИО некорректно
                    surname = full_name_parts[0]
                    name = full_name_parts[1]
                    patronymic = full_name_parts[2] if len(full_name_parts) > 2 else None

                    # Метро
                    metro_station = row.iloc[column_mapping["Ближайшее метро от дома"]] if pd.notna(
                        row.iloc[column_mapping["Ближайшее метро от дома"]]
                    ) else ""

                    # Телефон
                    phone_number = row.iloc[column_mapping["Мобильный телефон"]] if pd.notna(
                        row.iloc[column_mapping["Мобильный телефон"]]
                    ) else ""

                    # База практики
                    organization_name = row.iloc[column_mapping["База практики"]] if pd.notna(
                        row.iloc[column_mapping["База практики"]]
                    ) else ""

                    # Находим организацию (если существует)
                    organization = None
                    if organization_name and isinstance(organization_name, str):
                        organization, _ = Organization.objects.get_or_create(
                            full_name=organization_name.strip(),
                            defaults={
                                "legal_address": "",
                                "actual_address": "",
                                "inn": "",
                                "kpp": "",
                                "ogrn": "",
                                "phone_number": "",
                                "email": "",
                            }
                        )

                    # Создаем практиканта
                    intern = Intern(
                        last_name=surname,
                        first_name=name,
                        middle_name=patronymic,
                        phone_number=phone_number,
                        metro_station=metro_station,
                        group=group,
                        organization=organization
                    )
                    intern.clean()  # Проверяем валидность данных
                    intern.save()

            messages.success(request, "Данные успешно загружены.")
            return redirect('interns_list')

        except Exception as e:
            messages.error(request, f"Ошибка при загрузке данных: {str(e)}")

    return render(request, 'doc/upload_interns.html')


def download_filled_document_supervisor(request, document_id, intern_id):
    # Получаем документ и студента
    document = get_object_or_404(Document, id=document_id)
    intern = get_object_or_404(Intern, id=intern_id)

    if not document.is_auto_fillable:
        return HttpResponse("Этот документ не поддерживает автоматическое заполнение", status=400)

    # Проверяем авторизацию пользователя
    if not request.session.get('email'):
        return HttpResponse("Пользователь не авторизован", status=401)

    # Получаем связанные данные
    organization = intern.organization
    college_supervisor = intern.college_supervisor
    org_supervisor = OrganizationSupervisor.objects.filter(organization=organization).first() if organization else None
    group = intern.group

    # Подготавливаем данные для замены
    replacements = {
        '{{student_last_name}}': intern.last_name or '',
        '{{student_first_name}}': intern.first_name or '',
        '{{student_middle_name}}': intern.middle_name or '',
        '{{student_group}}': group.name if group else 'Не указана',
        '{{college_supervisor_name}}': (
            f"{college_supervisor.last_name} {college_supervisor.first_name} {college_supervisor.middle_name or ''}"
            if college_supervisor else "Не указан"
        ),
        '{{org_supervisor_name}}': (
            f"{org_supervisor.last_name} {org_supervisor.first_name} {org_supervisor.middle_name or ''}"
            if org_supervisor else "Не указан"
        ),
        '{{org_supervisor_position}}': org_supervisor.position if org_supervisor else "Не указана",
        '{{org_name}}': organization.full_name if organization else "Не указана",
        '{{org_legal_address}}': organization.legal_address if organization else "Не указан",
    }

    try:
        # Открываем документ
        doc_path = os.path.join(settings.MEDIA_ROOT, str(document.file))
        doc = DocxDocument(doc_path)

        # Функция для сохранения стилей
        def get_run_styles(run):
            return {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font': {
                    'name': run.font.name,
                    'size': run.font.size,
                    'color': run.font.color.rgb if run.font.color else None
                }
            }

        # Функция для применения стилей
        def apply_styles(run, styles):
            run.bold = styles['bold']
            run.italic = styles['italic']
            run.underline = styles['underline']
            if styles['font']['name']:
                run.font.name = styles['font']['name']
            if styles['font']['size']:
                run.font.size = styles['font']['size']
            if styles['font']['color']:
                run.font.color.rgb = styles['font']['color']

        # Обрабатываем параграфы
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    # Сохраняем стили всех runs в параграфе
                    styles = [get_run_styles(run) for run in paragraph.runs]

                    # Заменяем текст
                    paragraph.text = paragraph.text.replace(key, value)

                    # Восстанавливаем стили
                    for i, run in enumerate(paragraph.runs):
                        if i < len(styles):
                            apply_styles(run, styles[i])

        # Обрабатываем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                # Сохраняем стили
                                styles = [get_run_styles(run) for run in paragraph.runs]

                                # Заменяем текст
                                paragraph.text = paragraph.text.replace(key, value)

                                # Восстанавливаем стили
                                for i, run in enumerate(paragraph.runs):
                                    if i < len(styles):
                                        apply_styles(run, styles[i])

        # Создаем временный файл
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f"{document.title}.docx"
        safe_filename = quote(filename)

        # Отправляем файл
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{safe_filename}"'
        return response

    except Exception as e:
        logger.error(f"Ошибка при обработке документа: {str(e)}")
        return HttpResponse(f"Ошибка при обработке документа: {str(e)}", status=500)


def interns_list(request):
    # Проверка авторизации
    if not request.session.get('email'):
        return redirect('authPage')

    # Получение объекта пользователя на основе email из сессии
    try:
        user = Account.objects.get(email=request.session['email'])
    except Account.DoesNotExist:
        # Если пользователь не найден, перенаправляем на страницу авторизации
        return redirect('authPage')

    # Получение списка организаций, ожидающих подтверждения
    organizations = Organization.objects.filter(is_registration_request=True, is_approved=False)

    # Получение списка практикантов
    interns = Intern.objects.all()

    # Если пользователь - руководитель практики, фильтруем по управляемым группам
    if user.role.name == 'Руководитель практики':
        managed_groups = user.managed_groups.all()
        interns = interns.filter(group__in=managed_groups)

    # Формирование контекста
    context = {
        "organizations": organizations,
        "interns": interns,
    }

    return render(request, 'doc/interns_base.html', context)


def organizer_index(request):
    if not request.session.get('email'):
        return redirect('authPage')

    try:
        user = Account.objects.get(email=request.session['email'])
    except Account.DoesNotExist:
        return redirect('authPage')

    organization = None
    supervisor = None

    if user.role.name == 'Организация':
        # Ищем организацию по email пользователя
        organization = Organization.objects.filter(email=user.email).first()
        if organization:
            supervisor = OrganizationSupervisor.objects.filter(organization=organization).first()

    interns = Intern.objects.all()

    context = {
        "interns": interns,
        "organization": organization,
        "supervisor": supervisor,
        "user": user,
    }

    return render(request, 'doc/organizer_index.html', context)


def intern_detail(request, intern_id):
    intern = get_object_or_404(Intern, id=intern_id)
    return render(request, 'doc/intern_detail.html', {'intern': intern})


def add_intern(request):
    if request.method == 'POST':
        form = InternForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('interns_list')
    else:
        form = InternForm()
    return render(request, 'doc/add_intern.html', {'form': form})


@csrf_exempt
def update_intern(request, intern_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            field = data.get('field')
            value = data.get('value')

            # Проверяем, что поле не является id или organization
            if field in ['id', 'organization']:
                return JsonResponse({'success': False, 'error': 'Редактирование этого поля запрещено.'})

            intern = Intern.objects.get(id=intern_id)

            if field == 'last_name':
                intern.last_name = value
            elif field == 'first_name':
                intern.first_name = value
            elif field == 'middle_name':
                intern.middle_name = value
            elif field == 'phone_number':
                intern.phone_number = value
            elif field == 'email':
                intern.email = value
            elif field == 'metro_station':
                intern.metro_station = value

            intern.save()
            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@csrf_exempt
def request_resume_access(request, intern_id):
    if request.method == 'POST':
        try:
            intern = Intern.objects.get(id=intern_id)
            if intern.request_resume_access():
                return JsonResponse({'success': True, 'message': 'Доступ к резюме разрешен.'})
            else:
                return JsonResponse({'success': False, 'message': 'Доступ к резюме уже был разрешен.'})
        except Intern.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Практикант не найден.'})
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})


def documents_page(request):
    if not request.session.get('email'):
        return redirect('authPage')  # Перенаправляем на страницу авторизации, если пользователь не авторизован

    # Получаем текущего пользователя
    user = Account.objects.get(email=request.session['email'])

    # Получаем документы пользователя
    documents = Document.objects.filter(uploaded_by=user)

    # Получаем все практики
    practices = Practice.objects.all()

    # Получаем все группы, которыми управляет пользователь (если он руководитель практики)
    if user.role.name == 'Руководитель практики':
        groups = user.managed_groups.all()
    else:
        groups = Group.objects.all()

    # Получаем всех студентов из этих групп
    interns = Intern.objects.filter(group__in=groups)

    # Получаем все специальности
    specialties = Specialty.objects.all()

    context = {
        'documents': documents,
        'practices': practices,
        'groups': groups,
        'specialties': specialties,
    }

    return render(request, 'doc/documents.html', context)


@csrf_exempt
def upload_document_ajax(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.uploaded_by = Account.objects.get(email=request.session.get('email'))
            document.save()
            return JsonResponse({'success': True, 'message': 'Документ успешно загружен.'})
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})


def download_filled_document(request, document_id):
    document = get_object_or_404(Document, id=document_id)

    if not document.is_auto_fillable:
        return HttpResponse("Этот документ не поддерживает автоматическое заполнение", status=400)

    email = request.session.get('email')
    if not email:
        return HttpResponse("Пользователь не авторизован", status=401)

    account = Account.objects.filter(email=email).first()
    if not account:
        return HttpResponse("Аккаунт не найден", status=404)

    intern = Intern.objects.filter(email=email).first()
    organization = intern.organization if intern else None
    college_supervisor = intern.college_supervisor if intern else None
    org_supervisor = OrganizationSupervisor.objects.filter(organization=organization).first() if organization else None

    replacements = {
        '{{student_last_name}}': account.surname,
        '{{student_first_name}}': account.name,
        '{{student_middle_name}}': account.patronymic or '',
        '{{student_group}}': intern.group.name if intern and intern.group else 'Не указана',
        '{{college_supervisor_name}}': (
            f"{college_supervisor.last_name} {college_supervisor.first_name} {college_supervisor.middle_name or ''}"
            if college_supervisor else "Не указан"
        ),
        '{{org_supervisor_name}}': (
            f"{org_supervisor.last_name} {org_supervisor.first_name} {org_supervisor.middle_name or ''}"
            if org_supervisor else "Не указан"
        ),
        '{{org_supervisor_position}}': org_supervisor.position if org_supervisor else "Не указана",
        '{{org_name}}': organization.full_name if organization else "Не указана",
        '{{org_legal_address}}': organization.legal_address if organization else "Не указан",
    }

    try:
        doc = DocxDocument(os.path.join(settings.MEDIA_ROOT, str(document.file)))

        # Установка стиля по умолчанию для всего документа
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    # Применяем форматирование к абзацу
                    paragraph.style = doc.styles['Normal']
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
                            # Применяем форматирование к ячейке
                            for paragraph in cell.paragraphs:
                                paragraph.style = doc.styles['Normal']
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(14)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
        return response

    except Exception as e:
        return HttpResponse(f"Ошибка при обработке документа: {str(e)}", status=500)


def delete_document_ajax(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    document.delete()
    return JsonResponse({'success': True, 'message': 'Документ успешно удален.'})


def get_groups(request):
    specialty_id = request.GET.get('specialty_id')
    if specialty_id:
        groups = Group.objects.filter(specialty_id=specialty_id).values('id', 'name')
        return JsonResponse({'groups': list(groups)})
    return JsonResponse({'groups': []})


def prakties(request):
    # Получение данных для всех разделов
    practices = Practice.objects.all().distinct()
    schedules = Schedule.objects.all()
    groups = Group.objects.all()
    specialties = Specialty.objects.all()

    # Если пользователь - руководитель практики, фильтруем по управляемым группам
    user = Account.objects.get(email=request.session['email'])
    if user.role.name == 'Руководитель практики':
        managed_groups = user.managed_groups.all()
        practices = practices.filter(groups__in=managed_groups)
        groups = groups.filter(id__in=managed_groups)

    context = {
        'practices': practices,
        'schedules': schedules,
        'groups': groups,
        'specialties': specialties,
    }
    return render(request, 'doc/prakties.html', context)


def add_practice(request):
    if request.method == 'POST':
        form = PracticeForm(request.POST)
        if form.is_valid():
            practice = form.save(commit=False)  # Создаем объект, но не сохраняем в базу
            practice.save()  # Сохраняем основную модель
            form.save_m2m()  # Сохраняем связи ManyToMany
            messages.success(request, 'Практика успешно добавлена.')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
    return redirect('prakties')


def add_schedule(request):
    if request.method == 'POST':
        form = ScheduleForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'График успешно добавлен.')
        else:
            messages.error(request, 'Ошибка при добавлении графика.')
    return redirect('prakties')


def edit_schedule(request, schedule_id):
    schedule = get_object_or_404(Schedule, id=schedule_id)
    if request.method == 'POST':
        form = ScheduleForm(request.POST, instance=schedule)
        if form.is_valid():
            form.save()
            messages.success(request, 'График успешно обновлен.')
            return redirect('prakties')
        else:
            messages.error(request, 'Ошибка при обновлении графика.')
    else:
        form = ScheduleForm(instance=schedule)
    return render(request, 'doc/edit_schedule.html', {'form': form, 'schedule': schedule})


def delete_schedule(request, schedule_id):
    schedule = get_object_or_404(Schedule, id=schedule_id)
    if request.method == 'POST':
        schedule.delete()
        messages.success(request, 'График успешно удален.')
        return redirect('prakties')
    return render(request, 'doc/confirm_delete_schedule.html', {'schedule': schedule})


def add_group(request):
    if request.method == 'POST':
        form = GroupForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Группа успешно добавлена.')
        else:
            messages.error(request, 'Ошибка при добавлении группы.')
    return redirect('prakties')


def edit_group(request, group_id):
    group = get_object_or_404(Group, id=group_id)
    if request.method == 'POST':
        form = GroupForm(request.POST, instance=group)
        if form.is_valid():
            form.save()
            messages.success(request, 'Группа успешно обновлена.')
            return redirect('prakties')
        else:
            messages.error(request, 'Ошибка при обновлении группы.')
    else:
        form = GroupForm(instance=group)
    return render(request, 'doc/edit_group.html', {'form': form, 'group': group})


def delete_group(request, group_id):
    group = get_object_or_404(Group, id=group_id)
    if request.method == 'POST':
        group.delete()
        messages.success(request, 'Группа успешно удалена.')
        return redirect('prakties')
    return render(request, 'doc/confirm_delete_group.html', {'group': group})


def edit_practice(request, practice_id):
    practice = get_object_or_404(Practice, id=practice_id)
    if request.method == 'POST':
        form = PracticeForm(request.POST, instance=practice)
        if form.is_valid():
            form.save()
            messages.success(request, 'Практика успешно обновлена.')
            return redirect('prakties')
        else:
            messages.error(request, 'Ошибка при обновлении практики.')
    else:
        form = PracticeForm(instance=practice)
    return render(request, 'doc/edit_practice.html', {'form': form, 'practice': practice})


def delete_practice(request, practice_id):
    practice = get_object_or_404(Practice, id=practice_id)
    if request.method == 'POST':
        practice.delete()
        messages.success(request, 'Практика успешно удалена.')
        return redirect('prakties')
    return render(request, 'doc/confirm_delete_practice.html', {'practice': practice})


# Главная страница
def index(request):
    if not request.session.get('email'):
        return redirect('authPage')  # Перенаправляем на авторизацию, если пользователь не авторизован

    pending_registrations_count = Organization.objects.filter(is_registration_request=True, is_approved=False).count()
    context = {
        "pending_registrations_count": pending_registrations_count,
    }

    if request.session.get("role") == "Организация":
        organization_id = request.session.get("organization_id")
        organization = Organization.objects.get(id=organization_id)
        context["organization"] = organization

    return render(request, "doc/index.html", context)


@csrf_exempt
def register_organization(request):
    if request.method == "POST":
        form = OrganizationRegistrationForm(request.POST)
        if form.is_valid():
            organization = form.save(commit=False)
            organization.is_registration_request = True
            organization.save()
            messages.success(request, "Заявка на регистрацию организации успешно отправлена. Ожидайте подтверждения.")
            return redirect("authPage")
        else:
            for error in form.errors.values():
                messages.error(request, error)
    else:
        form = OrganizationRegistrationForm()
    return render(request, "doc/register_organization.html", {"form": form})


def organization_login(request):
    if request.method == "POST":
        form = OrganizationLoginForm(request.POST)
        if form.is_valid():
            organization = form.cleaned_data["organization"]
            # Сохраняем данные организации в сессии
            request.session["organization_id"] = organization.id
            request.session["role"] = "Организация"
            request.session["organization_name"] = organization.full_name
            messages.success(request, f"Вы успешно авторизовались как {organization.full_name}.")
            return redirect("index")
        else:
            for error in form.errors.values():
                messages.error(request, error)
    else:
        form = OrganizationLoginForm()
    return render(request, "doc/organization_login.html", {"form": form})


def organization_logout(request):
    if request.session.get("role") == "Организация":
        request.session.flush()
    return redirect("index")


def organizations_list(request):
    # Получаем только те организации, которые подали заявку на регистрацию и не подтверждены
    pending_organizations = Organization.objects.filter(is_registration_request=True, is_approved=False)
    return render(request, 'doc/organizations_list.html', {'organizations': pending_organizations})


logger = logging.getLogger(__name__)


@csrf_exempt
def approve_organization(request, organization_id):
    if request.method == 'POST':
        organization = Organization.objects.get(id=organization_id)

        if not organization.is_approved:
            organization.is_approved = True
            organization.is_registration_request = False
            organization.save()

            if organization.email:
                # Проверяем, существует ли уже аккаунт с таким email
                if Account.objects.filter(email=organization.email).exists():
                    return JsonResponse(
                        {'success': False, 'message': 'Аккаунт с таким email уже существует.'}
                    )

                # Создаем аккаунт для организации
                account = Account.objects.create(
                    email=organization.email,
                    surname=organization.full_name,
                    name="Organization",
                    role=Role.objects.get(name='Организация'),
                    password=organization.password  # Используем пароль, заданный при регистрации
                )

                # Отправляем письмо с подтверждением регистрации
                subject = 'Ваша регистрация подтверждена'
                message = f'Ваша организация "{organization.full_name}" успешно зарегистрирована на платформе.'
                from_email = settings.EMAIL_HOST_USER
                recipient_list = [organization.email]

                try:
                    send_mail(subject, message, from_email, recipient_list, fail_silently=False)
                    logger.info(f"Письмо успешно отправлено на {organization.email}")
                except Exception as e:
                    logger.error(f"Ошибка при отправке письма: {e}")

            return JsonResponse(
                {'success': True, 'message': f'Организация "{organization.full_name}" успешно подтверждена.'})
        else:
            return JsonResponse({'success': False, 'message': 'Организация уже подтверждена.'})
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})


def organization_detail(request, organization_id):
    # Проверка авторизации
    if not request.session.get('email'):
        return redirect('authPage')

    # Получение организации
    organization = get_object_or_404(Organization, id=organization_id)

    # Получение руководителя организации
    supervisor = OrganizationSupervisor.objects.filter(organization=organization).first()

    # Получение студентов, связанных с организацией
    interns = Intern.objects.filter(organization=organization)

    # Получаем все группы, связанные с этой организацией через студентов
    groups = Group.objects.filter(intern__organization=organization).distinct()

    # Для каждой группы получаем связанные практики (только ПМ или преддипломные)
    groups_data = []
    for group in groups:
        practices = Practice.objects.filter(
            groups=group
        ).filter(
            models.Q(pm__isnull=False) | models.Q(preddiplom=True)
        )
        groups_data.append({
            'group': group,
            'practices': practices
        })

    # Проверка прав доступа
    user = Account.objects.get(email=request.session['email'])
    is_organization_admin = (user.role.name == 'Организация' and
                             request.session.get('organization_id') == organization.id)
    is_admin = user.role.name == 'Администратор'

    if not (is_organization_admin or is_admin):
        messages.error(request, 'У вас нет прав для просмотра этой страницы')
        return redirect('index')

    context = {
        'organization': organization,
        'supervisor': supervisor,
        'interns': interns,
        'groups_data': groups_data,  # Добавляем данные о группах и практиках
        'can_edit': is_organization_admin or is_admin,
    }

    return render(request, 'doc/organization_detail.html', context)


@csrf_exempt
def update_organization(request, organization_id):
    if request.method != 'POST':
        return JsonResponse({
            'success': False,
            'error': 'Метод не разрешен. Используйте POST.'
        }, status=405)

    try:
        # Проверка авторизации
        if not request.session.get('email'):
            return JsonResponse({
                'success': False,
                'error': 'Требуется авторизация'
            }, status=401)

        # Получаем пользователя и организацию
        user = Account.objects.get(email=request.session['email'])
        organization = get_object_or_404(Organization, id=organization_id)

        # Упрощенная проверка прав доступа
        # Проверяем, что пользователь - организация и его email совпадает с email организации
        if user.role.name != 'Организация' or user.email != organization.email:
            return JsonResponse({
                'success': False,
                'error': 'Недостаточно прав для редактирования этой организации'
            }, status=403)

        # Парсинг JSON данных
        try:
            data = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse({
                'success': False,
                'error': 'Неверный формат JSON данных'
            }, status=400)

        # Обновление данных организации
        org_fields = {
            'full_name': data.get('full_name'),
            'legal_address': data.get('legal_address'),
            'actual_address': data.get('actual_address'),
            'inn': data.get('inn'),
            'kpp': data.get('kpp'),
            'ogrn': data.get('ogrn'),
            'phone_number': data.get('phone_number'),
            'email': data.get('email')
        }

        for field, value in org_fields.items():
            if value is not None:
                setattr(organization, field, value)

        try:
            organization.full_clean()
            organization.save()
        except ValidationError as e:
            return JsonResponse({
                'success': False,
                'error': 'Ошибка валидации данных организации',
                'details': dict(e)
            }, status=400)

        # Обновление данных руководителя
        supervisor, created = OrganizationSupervisor.objects.get_or_create(
            organization=organization
        )

        supervisor_fields = {
            'last_name': data.get('supervisor_last_name'),
            'first_name': data.get('supervisor_first_name'),
            'middle_name': data.get('supervisor_middle_name'),
            'phone_number': data.get('supervisor_phone'),
            'position': data.get('supervisor_position')
        }

        for field, value in supervisor_fields.items():
            if value is not None:
                setattr(supervisor, field, value)

        try:
            supervisor.full_clean()
            supervisor.save()
        except ValidationError as e:
            return JsonResponse({
                'success': False,
                'error': 'Ошибка валидации данных руководителя',
                'details': dict(e)
            }, status=400)

        return JsonResponse({
            'success': True,
            'message': 'Данные успешно обновлены'
        })

    except Account.DoesNotExist:
        return JsonResponse({
            'success': False,
            'error': 'Пользователь не найден'
        }, status=404)
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': 'Внутренняя ошибка сервера',
            'details': str(e)
        }, status=500)


@csrf_exempt
def update_supervisor(request, supervisor_id):
    if request.method == 'POST':
        try:
            # Проверка авторизации
            if not request.session.get('email'):
                return JsonResponse({'success': False, 'error': 'Не авторизован'})

            user = Account.objects.get(email=request.session['email'])
            supervisor = get_object_or_404(OrganizationSupervisor, id=supervisor_id)

            # Проверка прав
            is_organization_admin = (user.role.name == 'Организация' and
                                     request.session.get('organization_id') == supervisor.organization.id)
            is_admin = user.role.name == 'Администратор'

            if not (is_organization_admin or is_admin):
                return JsonResponse({'success': False, 'error': 'Нет прав для редактирования'})

            data = json.loads(request.body)
            field = data.get('field')
            value = data.get('value')

            # Запрещаем редактирование некоторых полей
            if field in ['id', 'organization']:
                return JsonResponse({'success': False, 'error': 'Редактирование этого поля запрещено'})

            # Обновляем поле
            if hasattr(supervisor, field):
                setattr(supervisor, field, value)
                supervisor.save()
                return JsonResponse({'success': True})
            else:
                return JsonResponse({'success': False, 'error': 'Неверное поле'})

        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

    return JsonResponse({'success': False, 'error': 'Неверный метод запроса'})


def edit_organization(request, organization_id):
    # Проверка авторизации
    if not request.session.get('email'):
        return redirect('authPage')

    user = Account.objects.get(email=request.session['email'])
    organization = get_object_or_404(Organization, id=organization_id)

    # Проверка прав
    is_organization_admin = (user.role.name == 'Организация' and
                             request.session.get('organization_id') == organization.id)
    is_admin = user.role.name == 'Администратор'

    if not (is_organization_admin or is_admin):
        messages.error(request, 'У вас нет прав для редактирования этой организации')
        return redirect('index')

    supervisor = OrganizationSupervisor.objects.filter(organization=organization).first()

    if request.method == 'POST':
        org_form = OrganizationForm(request.POST, instance=organization)
        supervisor_form = OrganizationSupervisorForm(request.POST, instance=supervisor)

        if org_form.is_valid() and supervisor_form.is_valid():
            org_form.save()
            sup = supervisor_form.save(commit=False)
            sup.organization = organization
            sup.save()
            messages.success(request, 'Данные успешно обновлены')
            return redirect('organization_detail', organization_id=organization.id)
    else:
        org_form = OrganizationForm(instance=organization)
        supervisor_form = OrganizationSupervisorForm(instance=supervisor)

    context = {
        'organization': organization,
        'org_form': org_form,
        'supervisor_form': supervisor_form,
    }

    return render(request, 'doc/edit_organization.html', context)


def student_index(request):
    return render(request, 'doc/student_index.html')


@csrf_exempt
def update_intern_skills(request):
    if request.method == 'POST':
        try:
            intern_id = request.POST.get('intern_id')
            existing_tags = json.loads(request.POST.get('existing_tags', '[]'))
            new_tags_text = request.POST.get('new_tags', '')

            intern = Intern.objects.get(id=intern_id)

            # Очищаем текущие теги
            intern.tags.clear()

            # Добавляем выбранные существующие теги
            for tag_id in existing_tags:
                tag = Tag.objects.get(id=tag_id)
                intern.tags.add(tag)

            # Добавляем новые теги
            new_tags_list = [tag.strip() for tag in new_tags_text.split(',') if tag.strip()]
            for tag_name in new_tags_list:
                tag, created = Tag.objects.get_or_create(name=tag_name)
                intern.tags.add(tag)

            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


@csrf_exempt
def upload_intern_resume(request):
    if request.method == 'POST':
        try:
            intern_id = request.POST.get('intern_id')
            intern = Intern.objects.get(id=intern_id)

            # Обработка удаления резюме
            if 'delete_resume' in request.POST:
                if intern.resume:
                    intern.resume.delete()
                intern.resume = None
                intern.save()
                return JsonResponse({'success': True})

            # Обработка загрузки нового резюме
            if 'resume_file' in request.FILES:
                # Удаляем старое резюме, если оно есть
                if intern.resume:
                    intern.resume.delete()

                # Сохраняем новое резюме
                intern.resume = request.FILES['resume_file']
                intern.save()
                return JsonResponse({'success': True})

            return JsonResponse({'success': False, 'error': 'No file provided'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})


# Вывод данных пользователя из сессии
def account(request):
    if not request.session.get('email'):
        return redirect('authPage')

    user = Account.objects.get(email=request.session.get('email'))
    intern = Intern.objects.filter(email=user.email).first()
    all_tags = Tag.objects.all()  # Получаем все существующие теги

    if intern:
        documents = Document.objects.filter(practice__groups=intern.group)
        org_supervisor = OrganizationSupervisor.objects.filter(organization=intern.organization).first()
    else:
        documents = []
        org_supervisor = None

    context = {
        'intern': intern,
        'documents': documents,
        'org_supervisor': org_supervisor,
        'all_tags': all_tags,  # Передаем все теги в шаблон
    }
    return render(request, 'doc/account.html', context)


@csrf_exempt
def change_student_email(request):
    if request.method == 'POST':
        try:
            current_email = request.POST.get('current_email')
            new_email = request.POST.get('new_email')
            password = request.POST.get('confirm_password')

            # Проверяем, что пользователь авторизован
            if not request.session.get('email') or request.session.get('email') != current_email:
                return JsonResponse({'success': False, 'error': 'Недостаточно прав для изменения email.'})

            # Получаем аккаунт студента
            account = Account.objects.get(email=current_email)

            # Проверяем пароль
            if not account.check_password(password):
                return JsonResponse({'success': False, 'error': 'Неверный пароль.'})

            # Проверяем, что новый email не занят
            if Account.objects.filter(email=new_email).exists():
                return JsonResponse({'success': False, 'error': 'Этот email уже используется.'})

            # Обновляем email в аккаунте
            account.email = new_email
            account.save()

            # Обновляем email в модели Intern, если она существует
            intern = Intern.objects.filter(email=current_email).first()
            if intern:
                intern.email = new_email
                intern.save()

            return JsonResponse({'success': True})
        except Account.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Аккаунт не найден.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Недопустимый метод запроса.'})


@csrf_exempt
def send_password(request, intern_id):
    if request.method == 'POST':
        try:
            logger.info(f"Attempting to send password for intern_id: {intern_id}")
            intern = Intern.objects.get(id=intern_id)
            if intern.email:
                account = Account.objects.get(email=intern.email)
                account.generate_and_send_password()
                return JsonResponse({'success': True, 'message': 'Пароль успешно отправлен.'})
            else:
                return JsonResponse({'success': False, 'message': 'У студента не указан email.'})
        except Intern.DoesNotExist:
            logger.error(f"Intern with id {intern_id} does not exist.")
            return JsonResponse({'success': False, 'message': 'Студент не найден.'})
        except Account.DoesNotExist:
            logger.error(f"Account for intern_id {intern_id} does not exist.")
            return JsonResponse({'success': False, 'message': 'Аккаунт студента не найден.'})
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})


@csrf_exempt
def send_passwords_to_all_students(request):
    if request.method == 'POST':
        try:
            students = Account.objects.filter(role__name='Студент', email_sent=False)
            for student in students:
                student.generate_and_send_password()
            return JsonResponse({'success': True, 'message': 'Пароли успешно отправлены всем студентам.'})
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)})
    return JsonResponse({'success': False, 'message': 'Недопустимый метод запроса.'})


@csrf_exempt
def send_interview_invitation(request):
    if request.method != 'POST':
        return JsonResponse({'success': False, 'message': 'Метод не разрешен'}, status=405)

    try:
        # Проверка авторизации организации
        if not request.session.get('role') == 'Организация':
            return JsonResponse(
                {'success': False, 'message': 'Только организации могут отправлять приглашения'},
                status=403
            )

        data = json.loads(request.body)

        # Проверяем обязательные поля
        required_fields = ['intern_id', 'interview_date', 'interview_location']
        for field in required_fields:
            if field not in data:
                return JsonResponse(
                    {'success': False, 'message': f'Отсутствует обязательное поле: {field}'},
                    status=400
                )

        # Получаем данные студента
        intern = Intern.objects.get(id=data['intern_id'])
        if not intern.email:
            return JsonResponse(
                {'success': False, 'message': 'У студента не указан email'},
                status=400
            )

        # Получаем аккаунт организации
        org_account = Account.objects.get(email=request.session['email'])

        # Проверяем валидность email-адресов
        from django.core.validators import validate_email
        from django.core.exceptions import ValidationError

        try:
            validate_email(intern.email)
            validate_email(org_account.email)
        except ValidationError as e:
            return JsonResponse(
                {'success': False, 'message': f'Некорректный email-адрес: {str(e)}'},
                status=400
            )

        # Преобразуем строку даты в datetime с часовым поясом
        from django.utils.timezone import make_aware
        from datetime import datetime
        naive_datetime = datetime.strptime(data['interview_date'], '%Y-%m-%dT%H:%M')
        aware_datetime = make_aware(naive_datetime)

        # Создаем приглашение
        invitation = InterviewInvitation.objects.create(
            intern=intern,
            interview_date=aware_datetime,
            location=data['interview_location'],
            message=data.get('interview_message', ''),
            created_by=org_account
        )

        # Форматируем дату для письма
        formatted_date = aware_datetime.strftime('%d.%m.%Y в %H:%M')

        # Отправляем письмо
        context = {
            'intern': intern,
            'invitation': invitation,
            'formatted_date': formatted_date,
            'sender_email': settings.EMAIL_HOST_USER,  # Используем EMAIL_HOST_USER
            'sender_name': f"{org_account.surname} {org_account.name}"
        }

        text_content = render_to_string('emails/interview_invitation.txt', context)
        html_content = render_to_string('emails/interview_invitation.html', context)

        email = EmailMultiAlternatives(
            subject=f'Приглашение на собеседование от {context["sender_name"]}',
            body=text_content,
            from_email=settings.EMAIL_HOST_USER,  # Отправляем с EMAIL_HOST_USER
            to=[intern.email],
            reply_to=[org_account.email]  # Для ответов используем email организации
        )
        email.attach_alternative(html_content, "text/html")
        email.send(fail_silently=False)

        return JsonResponse({
            'success': True,
            'message': 'Приглашение успешно отправлено',
            'invitation_id': invitation.id
        })

    except Intern.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Студент не найден'}, status=404)
    except Account.DoesNotExist:
        return JsonResponse({'success': False, 'message': 'Аккаунт организации не найден'}, status=404)
    except Exception as e:
        logger.error(f"Ошибка при отправке приглашения: {str(e)}", exc_info=True)
        return JsonResponse(
            {'success': False, 'message': f'Ошибка при отправке приглашения: {str(e)}'},
            status=500
        )


@csrf_exempt
def change_password(request):
    if request.method == 'POST':
        if not request.session.get('email'):
            return JsonResponse({'success': False, 'error': 'Пользователь не авторизован.'})

        user = Account.objects.get(email=request.session['email'])
        old_password = request.POST.get('old_password')
        new_password = request.POST.get('new_password')
        confirm_new_password = request.POST.get('confirm_new_password')

        if not user.check_password(old_password):
            return JsonResponse({'success': False, 'error': 'Старый пароль введен неверно.'})

        if new_password != confirm_new_password:
            return JsonResponse({'success': False, 'error': 'Новый пароль и подтверждение не совпадают.'})

        user.set_password(new_password)
        user.save()
        return JsonResponse({'success': True})

    return JsonResponse({'success': False, 'error': 'Недопустимый метод запроса.'})


# Реализация выхода из сессии пользователя
def logout_view(request):
    if request.session.get('email'):
        request.session.flush()  # Очистка сессии
        messages.success(request, 'Вы успешно вышли из аккаунта.')
    else:
        messages.info(request, 'Вы уже не были в системе.')
    return redirect('authPage')


# Реализация функции регистрации пользователя и сохранение его в сессии
def register(request):
    if request.method == 'POST':
        form = UserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user_role = Role.objects.get(
                name='Руководитель практики')
            user.role = user_role
            user.save()

            # Сохраняем данные пользователя в сессию
            request.session['email'] = user.email
            request.session['role'] = user.role.name
            request.session.modified = True

            messages.success(request, 'Вы успешно зарегистрировались!')
            return redirect('indexPage')
        else:
            messages.error(request, 'Форма содержит ошибки.')
    else:
        form = UserRegisterForm()
    return render(request, 'doc/reg.html', {'form': form})


def auth(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            try:
                # Поиск пользователя по email
                user = Account.objects.get(email=email)

                # Проверка пароля
                if user.check_password(password):
                    # Сохраняем данные пользователя в сессии
                    request.session['email'] = user.email
                    request.session['role'] = user.role.name
                    request.session['user_surname'] = user.surname
                    request.session['user_name'] = user.name
                    request.session['user_patronymic'] = user.patronymic
                    request.session.modified = True

                    # Перенаправление в зависимости от роли
                    if user.role.name == "Руководитель практики":
                        return redirect('indexPage')  # Перенаправление на главную страницу
                    elif user.role.name == "Организация":
                        return redirect('organizations_index')
                    elif user.role.name == "Администратор":
                        return redirect('admin_panel')  # Перенаправление на панель администратора
                    elif user.role.name == "Студент":
                        return redirect('useraccountPage')  # Перенаправление на страницу студента
                    else:
                        messages.error(request, 'Неизвестная роль пользователя.')
                        return redirect('auth')  # Возвращаем обратно на страницу авторизации
                else:
                    messages.error(request, 'Неверный email или пароль.')
            except Account.DoesNotExist:
                messages.error(request, 'Пользователь с таким email не найден.')
        else:
            messages.error(request, 'Некорректные данные в форме.')
    else:
        form = LoginForm()

    return render(request, 'doc/auth.html', {'form': form})


def admin_panel(request):
    # Получаем данные для всех моделей с пагинацией
    accounts = Account.objects.all().order_by('surname')
    interns = Intern.objects.all().order_by('last_name')
    groups = Group.objects.all().order_by('name')  # Добавлен order_by
    organizations = Organization.objects.all().order_by('full_name')  # Добавлен order_by
    college_supervisors = CollegeSupervisor.objects.all().order_by('last_name')  # Добавлен order_by
    specialties = Specialty.objects.all().order_by('name')  # Добавлен order_by
    org_supervisors = OrganizationSupervisor.objects.all().order_by('last_name')  # Добавлен order_by
    roles = Role.objects.all().order_by('name')  # Добавлен order_by
    tags = Tag.objects.all().order_by('name')  # Добавлен order_by
    schedules = Schedule.objects.all().order_by('schedule_description')  # Добавлен order_by (сортировка по дате)
    practices = Practice.objects.all().order_by('pp')  # Добавлен order_by
    documents = Document.objects.all().order_by('-uploaded_at')  # Добавлен order_by
    students = Student.objects.all().order_by('account')  # Добавлен order_by
    document_links = DocumentLinks.objects.all().order_by('document_link')  # Добавлен order_by

    # Фильтрация аккаунтов по роли
    selected_role = request.GET.get('role')
    if selected_role:
        accounts = accounts.filter(role_id=selected_role)

    # Фильтрация студентов по группе
    selected_group = request.GET.get('group')
    if selected_group:
        interns = interns.filter(group_id=selected_group)

    # Пагинация для всех моделей
    paginator_accounts = Paginator(accounts, 30)
    paginator_interns = Paginator(interns, 30)
    paginator_groups = Paginator(groups, 20)
    paginator_organizations = Paginator(organizations, 20)
    paginator_college_supervisors = Paginator(college_supervisors, 20)
    paginator_specialties = Paginator(specialties, 20)
    paginator_org_supervisors = Paginator(org_supervisors, 20)
    paginator_roles = Paginator(roles, 10)
    paginator_tags = Paginator(tags, 10)
    paginator_schedules = Paginator(schedules, 20)
    paginator_practices = Paginator(practices, 20)
    paginator_documents = Paginator(documents, 20)
    paginator_students = Paginator(students, 20)
    paginator_document_links = Paginator(document_links, 20)

    page_number = request.GET.get('page')
    accounts_page = paginator_accounts.get_page(page_number)
    interns_page = paginator_interns.get_page(page_number)
    groups_page = paginator_groups.get_page(page_number)
    organizations_page = paginator_organizations.get_page(page_number)
    college_supervisors_page = paginator_college_supervisors.get_page(page_number)
    specialties_page = paginator_specialties.get_page(page_number)
    org_supervisors_page = paginator_org_supervisors.get_page(page_number)
    roles_page = paginator_roles.get_page(page_number)
    tags_page = paginator_tags.get_page(page_number)
    schedules_page = paginator_schedules.get_page(page_number)
    practices_page = paginator_practices.get_page(page_number)
    documents_page = paginator_documents.get_page(page_number)
    students_page = paginator_students.get_page(page_number)
    document_links_page = paginator_document_links.get_page(page_number)

    context = {
        'accounts': accounts_page,
        'interns': interns_page,
        'groups': groups_page,
        'organizations': organizations_page,
        'college_supervisors': college_supervisors_page,
        'specialties': specialties_page,
        'organization_supervisors': org_supervisors_page,
        'roles': roles_page,
        'tags': tags_page,
        'schedules': schedules_page,
        'practices': practices_page,
        'documents': documents_page,
        'students': students_page,
        'document_links': document_links_page,
        'all_groups': groups,  # Для выпадающего списка групп
        'all_roles': roles,  # Для выпадающего списка ролей
    }
    return render(request, 'doc/admin_panel.html', context)


def admin_add(request, model_name):
    try:
        # Получаем класс формы для указанной модели
        form_class = get_form_by_model_name(model_name)
    except ValueError as e:
        raise Http404(str(e))  # Возвращаем 404, если модель не найдена

    if request.method == 'POST':
        # Создаем экземпляр формы с переданными данными
        form = form_class(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, f'Запись успешно добавлена.')
            return redirect('admin_panel')
        else:
            messages.error(request, 'Ошибка при добавлении записи.')
    else:
        # Создаем пустую форму для GET-запроса
        form = form_class()

    # Подготавливаем данные о полях формы для передачи в шаблон
    form_fields = []
    for field_name, field in form.fields.items():
        field_data = {
            'field': form[field_name],  # Объект поля формы
            'related_model_name': None,  # По умолчанию нет связанной модели
        }

        # Проверяем, является ли поле связанным с другой моделью
        if hasattr(form.Meta.model, field_name) and hasattr(getattr(form.Meta.model, field_name), 'field'):
            related_model = getattr(form.Meta.model, field_name).field.related_model
            if related_model:
                field_data['related_model_name'] = related_model._meta.model_name

        form_fields.append(field_data)

    return render(request, 'doc/admin_add.html', {
        'form': form,
        'form_fields': form_fields,  # Передаем данные о полях формы
        'model_name': model_name,
    })


def admin_edit(request, model_name, pk):
    model = get_model_by_name(model_name)
    instance = get_object_or_404(model, id=pk)

    if request.method == 'POST':
        form = get_form_by_model_name(model_name)(request.POST, request.FILES, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, f'Запись успешно обновлена.')
            return redirect('admin_panel')
        else:
            messages.error(request, 'Ошибка при обновлении записи.')
    else:
        form = get_form_by_model_name(model_name)(instance=instance)

    return render(request, 'doc/admin_edit.html', {'form': form, 'model_name': model_name, 'pk': pk})


def admin_delete(request, model_name, pk):
    model = get_model_by_name(model_name)
    instance = get_object_or_404(model, id=pk)

    if request.method == 'POST':
        instance.delete()
        messages.success(request, f'Запись успешно удалена.')
        return redirect('admin_panel')

    return render(request, 'doc/admin_delete.html', {'model_name': model_name, 'pk': pk})


# Вспомогательные функции
def get_model_by_name(model_name):
    """
    Возвращает класс модели по её имени.

    :param model_name: Имя модели (строка).
    :return: Класс модели.
    :raises ValueError: Если модель с таким именем не найдена.
    """
    models_dict = {
        'intern': Intern,
        'group': Group,
        'organization': Organization,
        'collegesupervisor': CollegeSupervisor,  # Обратите внимание на имя модели
        'specialty': Specialty,
        'organizationsupervisor': OrganizationSupervisor,  # Обратите внимание на имя модели
        'role': Role,
        'tag': Tag,
        'schedule': Schedule,
        'practice': Practice,
        'account': Account,
        'document': Document,
        'student': Student,
        'documentlinks': DocumentLinks,
    }
    model_class = models_dict.get(model_name.lower())  # Приводим к нижнему регистру для надежности
    if model_class is None:
        raise ValueError(f"Модель с именем '{model_name}' не найдена.")
    return model_class


def get_form_by_model_name(model_name, *args, **kwargs):
    """
    Возвращает класс формы для модели по её имени.

    :param model_name: Имя модели (строка).
    :param args: Аргументы для передачи в форму.
    :param kwargs: Ключевые аргументы для передачи в форму.
    :return: Класс формы.
    :raises ValueError: Если форма для модели не найдена.
    """
    forms_dict = {
        'intern': InternForm,
        'group': GroupForm,
        'organization': OrganizationForm,
        'collegesupervisor': CollegeSupervisorForm,  # Обратите внимание на имя модели
        'specialty': SpecialtyForm,
        'organizationsupervisor': OrganizationSupervisorForm,  # Обратите внимание на имя модели
        'role': RoleForm,
        'tag': TagForm,
        'schedule': ScheduleForm,
        'practice': PracticeForm,
        'account': AccountForm,
        'document': DocumentForm,
        'student': StudentForm,
        'documentlinks': DocumentLinksForm,
    }
    form_class = forms_dict.get(model_name.lower())  # Приводим к нижнему регистру для надежности
    if form_class is None:
        raise ValueError(f"Форма для модели '{model_name}' не найдена.")
    return form_class
