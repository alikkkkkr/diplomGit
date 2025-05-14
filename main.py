import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx.enum.section import WD_ORIENTATION
from ttkthemes import ThemedStyle
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from settings_loader import *
from doc.models import *
from datetime import datetime
import traceback
from docx import Document
import os
import subprocess
from django.core.files import File
from django.conf import settings


class StudentManagementApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Управление студентами")
        self.root.geometry("1200x700")
        self.root.configure(bg="white")

        # Инициализация стилей
        self.setup_styles()

        # Загрузка данных
        self.students_data = self.fetch_students()

        # Создание интерфейса
        self.create_widgets()

        # Обновление таблицы
        self.update_table(self.students_data)

    def setup_styles(self):
        """Настройка стилей интерфейса"""
        self.style = ThemedStyle(self.root)
        self.style.set_theme("arc")

        # Настройка цветовой палитры
        self.style.configure("TLabel", font=("Oswald", 14), foreground="#333", background="white")
        self.style.configure("TButton", font=("Oswald", 12, "bold"), background="#f8f9fa",
                             foreground="#333", padding=10, borderwidth=1)
        self.style.map("TButton", background=[("active", "#e9ecef")])
        self.style.configure("TEntry", font=("Oswald", 12), foreground="#333", background="white", padding=5)
        self.style.configure("Treeview", font=("Oswald", 12), rowheight=30, background="white",
                             foreground="#333", fieldbackground="white")
        self.style.configure("Treeview.Heading", font=("Oswald", 14, "bold"), background="#f8f9fa",
                             foreground="#333")
        self.style.map("Treeview", background=[("selected", "#E0E0E0")])

    def fetch_students(self):
        """Получение списка студентов из базы данных"""
        try:
            interns = Intern.objects.select_related('group', 'organization').all()
            data = [
                (
                    intern.id,
                    intern.last_name,
                    intern.first_name,
                    intern.middle_name or "",
                    intern.phone_number or "",
                    intern.metro_station or "",
                    intern.group.name if intern.group else "",
                    intern.organization.full_name if intern.organization else "",
                    intern.college_supervisor.last_name if intern.college_supervisor else ""
                )
                for intern in interns
            ]
            return data
        except Exception as e:
            self.show_error("Ошибка при получении данных", str(e))
            return []

    def create_widgets(self):
        """Создание элементов интерфейса"""
        # Верхняя панель
        self.header_frame = ttk.Frame(self.root)
        self.header_frame.pack(pady=20, padx=20, fill=tk.X)

        self.title_label = ttk.Label(self.header_frame, text="Управление студентами", font=("Oswald", 24, "bold"))
        self.title_label.pack(side=tk.LEFT)

        self.button_frame = ttk.Frame(self.header_frame)
        self.button_frame.pack(side=tk.RIGHT)

        # Кнопки действий
        self.upload_btn = ttk.Button(self.button_frame, text="Загрузить Excel",
                                     command=self.upload_excel_file)
        self.upload_btn.pack(side=tk.LEFT, padx=5)

        self.add_btn = ttk.Button(self.button_frame, text="Добавить",
                                  command=self.open_add_student_modal)
        self.add_btn.pack(side=tk.LEFT, padx=5)

        self.generate_btn = ttk.Button(self.button_frame, text="Создать отчет",
                                       command=self.open_generate_file_modal)
        self.generate_btn.pack(side=tk.LEFT, padx=5)

        # Добавляем кнопку управления бэкапами
        self.backup_btn = ttk.Button(self.button_frame, text="Бэкапы БД",
                                     command=self.open_backup_management_modal)
        self.backup_btn.pack(side=tk.LEFT, padx=5)

        # Панель поиска и фильтрации
        self.filter_frame = ttk.Frame(self.root)
        self.filter_frame.pack(pady=10, padx=20, fill=tk.X)

        self.search_label = ttk.Label(self.filter_frame, text="Поиск:")
        self.search_label.pack(side=tk.LEFT, padx=5)

        self.search_entry = ttk.Entry(self.filter_frame, width=40)
        self.search_entry.pack(side=tk.LEFT, padx=5)
        self.search_entry.bind("<Return>", lambda e: self.filter_students(self.search_entry.get()))

        self.search_btn = ttk.Button(self.filter_frame, text="Найти",
                                     command=lambda: self.filter_students(self.search_entry.get()))
        self.search_btn.pack(side=tk.LEFT, padx=5)

        self.group_label = ttk.Label(self.filter_frame, text="Группа:")
        self.group_label.pack(side=tk.LEFT, padx=10)

        self.groups = ["Все группы"] + list(Group.objects.values_list('name', flat=True))
        self.group_var = tk.StringVar(value="Все группы")
        self.group_dropdown = ttk.Combobox(self.filter_frame, textvariable=self.group_var,
                                           values=self.groups, state="readonly", width=20)
        self.group_dropdown.pack(side=tk.LEFT, padx=5)

        self.filter_btn = ttk.Button(self.filter_frame, text="Фильтр",
                                     command=lambda: self.filter_by_group(self.group_var.get()))
        self.filter_btn.pack(side=tk.LEFT, padx=5)

        # Таблица студентов
        self.table_frame = ttk.Frame(self.root)
        self.table_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

        self.columns = ("ID", "Фамилия", "Имя", "Отчество", "Телефон", "Метро", "Группа", "Организация", "Руководитель")
        self.table = ttk.Treeview(self.table_frame, columns=self.columns, show="headings", height=20)

        # Настройка столбцов
        col_widths = [50, 150, 120, 120, 120, 100, 100, 200, 150]
        for col, width in zip(self.columns, col_widths):
            self.table.column(col, width=width, anchor=tk.W if col not in ["ID", "Телефон"] else tk.CENTER)
            self.table.heading(col, text=col)

        self.table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Скроллбар
        self.scrollbar = ttk.Scrollbar(self.table_frame, orient=tk.VERTICAL, command=self.table.yview)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.table.configure(yscrollcommand=self.scrollbar.set)

        # Контекстное меню
        self.setup_context_menu()

        # Привязка изменения размера
        self.root.bind("<Configure>", self.resize_table)

    def open_backup_management_modal(self):
        """Открытие модального окна для управления бэкапами"""
        self.backup_modal = tk.Toplevel(self.root)
        self.backup_modal.title("Управление бэкапами базы данных")
        self.backup_modal.geometry("900x650")
        self.backup_modal.resizable(False, False)
        self.backup_modal.grab_set()

        # Настройка стилей для модального окна
        style = ttk.Style(self.backup_modal)
        style.configure("Bold.TLabel", font=("Helvetica", 12, "bold"), foreground="#333333")

        # Стили для кнопок с черным текстом
        style.configure("Blue.TButton", font=("Helvetica", 10, "bold"),
                        foreground="black", background="#e6f2ff", padding=8, borderwidth=1)
        style.map("Blue.TButton",
                  background=[("active", "#cce0ff")])

        style.configure("Red.TButton", font=("Helvetica", 10, "bold"),
                        foreground="black", background="#ffebee", padding=8, borderwidth=1)
        style.map("Red.TButton",
                  background=[("active", "#ffcdd2")])

        style.configure("Gray.TButton", font=("Helvetica", 10, "bold"),
                        foreground="black", background="#f5f5f5", padding=8, borderwidth=1)
        style.map("Gray.TButton",
                  background=[("active", "#e0e0e0")])

        # Основной контейнер
        container = ttk.Frame(self.backup_modal)
        container.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)

        # Заголовок
        header_frame = ttk.Frame(container)
        header_frame.pack(fill=tk.X, pady=(0, 15))

        ttk.Label(header_frame, text="Управление бэкапами базы данных",
                  style="Bold.TLabel").pack(side=tk.LEFT)

        # Кнопка создания бэкапа (светло-синий фон)
        create_btn = ttk.Button(header_frame, text="＋ Создать новый бэкап",
                                style="Blue.TButton", command=self.create_database_backup)
        create_btn.pack(side=tk.RIGHT, padx=5)

        # Таблица с бэкапами
        table_frame = ttk.Frame(container)
        table_frame.pack(fill=tk.BOTH, expand=True)

        # Создаем Treeview с полосой прокрутки
        columns = ("ID", "Дата создания", "Размер", "Файл")
        self.backup_table = ttk.Treeview(
            table_frame,
            columns=columns,
            show="headings",
            selectmode="browse",
            height=15
        )

        # Настройка столбцов
        self.backup_table.column("ID", width=50, anchor=tk.CENTER)
        self.backup_table.column("Дата создания", width=180, anchor=tk.CENTER)
        self.backup_table.column("Размер", width=100, anchor=tk.CENTER)
        self.backup_table.column("Файл", width=400, anchor=tk.W)

        # Заголовки столбцов
        for col in columns:
            self.backup_table.heading(col, text=col)

        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.backup_table.yview)
        self.backup_table.configure(yscrollcommand=scrollbar.set)

        # Размещаем таблицу и полосу прокрутки
        self.backup_table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Панель действий
        action_frame = ttk.Frame(container)
        action_frame.pack(fill=tk.X, pady=(15, 0))

        # Кнопка восстановления (светло-серый фон)
        restore_btn = ttk.Button(
            action_frame,
            text="↻ Восстановить выбранный",
            style="Gray.TButton",
            command=self.restore_selected_backup
        )
        restore_btn.pack(side=tk.LEFT, padx=5)

        # Кнопка удаления (светло-красный фон)
        delete_btn = ttk.Button(
            action_frame,
            text="✕ Удалить выбранный",
            style="Red.TButton",
            command=self.delete_selected_backup
        )
        delete_btn.pack(side=tk.LEFT, padx=5)

        # Кнопка закрытия (светло-серый фон)
        close_btn = ttk.Button(
            action_frame,
            text="Закрыть",
            style="Gray.TButton",
            command=self.backup_modal.destroy
        )
        close_btn.pack(side=tk.RIGHT, padx=5)

        # Заполняем таблицу бэкапами
        self.load_backups_list()

    def load_backups_list(self):
        """Загрузка списка бэкапов в таблицу"""
        try:
            # Очищаем таблицу
            for row in self.backup_table.get_children():
                self.backup_table.delete(row)

            # Получаем бэкапы из базы данных
            backups = DatabaseBackup.objects.all().order_by('-created_at')

            # Заполняем таблицу
            for backup in backups:
                self.backup_table.insert("", "end", values=(
                    backup.id,
                    backup.created_at.strftime('%d.%m.%Y %H:%M'),
                    backup.size,
                    backup.file.name.split('/')[-1]  # Только имя файла без пути
                ))

            # Автоматически выбираем первый элемент, если он есть
            if len(self.backup_table.get_children()) > 0:
                self.backup_table.selection_set(self.backup_table.get_children()[0])
        except Exception as e:
            self.show_error("Ошибка загрузки бэкапов", str(e))

    def create_database_backup(self):
        """Создание нового бэкапа базы данных"""
        try:
            # Создаем имя файла с временной меткой
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f"backup_{timestamp}.dump"

            # Создаем временную директорию внутри MEDIA_ROOT
            temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp_backups')
            os.makedirs(temp_dir, exist_ok=True)

            # Полный безопасный путь к файлу
            temp_path = os.path.join(temp_dir, backup_filename)

            # Команда для создания бэкапа
            db_settings = settings.DATABASES['default']
            command = (
                f"pg_dump -h {db_settings['HOST']} -p {db_settings['PORT']} "
                f"-U {db_settings['USER']} -F c -b -v -f {temp_path} {db_settings['NAME']}"
            )

            # Запускаем команду
            env = os.environ.copy()
            env['PGPASSWORD'] = db_settings['PASSWORD']
            result = subprocess.run(command, shell=True, check=True, env=env,
                                    stdout=subprocess.PIPE, stderr=subprocess.PIPE)

            if result.returncode != 0:
                raise Exception(f"Ошибка создания бэкапа: {result.stderr.decode()}")

            # Проверяем, что файл был создан
            if not os.path.exists(temp_path):
                raise Exception("Файл бэкапа не был создан")

            # Открываем файл в бинарном режиме
            with open(temp_path, 'rb') as f:
                # Создаем File объект с безопасным именем
                backup_file = File(f, name=f"database_backups/{backup_filename}")

                # Создаем запись о бэкапе
                backup = DatabaseBackup(
                    file=backup_file,
                    created_by=self.get_current_user_account()
                )
                backup.save()

            # Удаляем временный файл
            try:
                os.remove(temp_path)
            except Exception as e:
                print(f"Ошибка при удалении временного файла: {e}")

            # Обновляем список бэкапов
            self.load_backups_list()

            messagebox.showinfo("Успех", "Бэкап базы данных успешно создан")
        except subprocess.CalledProcessError as e:
            error_msg = e.stderr.decode() if e.stderr else str(e)
            self.show_error("Ошибка создания бэкапа", f"Команда завершилась с ошибкой:\n{error_msg}")
        except Exception as e:
            self.show_error("Ошибка создания бэкапа", str(e))

    def restore_selected_backup(self):
        """Восстановление базы данных из выбранного бэкапа"""
        try:
            selected_item = self.backup_table.selection()
            if not selected_item:
                messagebox.showwarning("Предупреждение", "Выберите бэкап для восстановления")
                return

            # Получаем ID выбранного бэкапа
            backup_id = self.backup_table.item(selected_item[0])['values'][0]
            backup = DatabaseBackup.objects.get(id=backup_id)

            # Подтверждение действия
            if not messagebox.askyesno(
                    "Подтверждение",
                    f"Вы уверены, что хотите восстановить базу данных из бэкапа от {backup.created_at.strftime('%d.%m.%Y %H:%M')}?\n"
                    "Это действие перезапишет текущие данные!"
            ):
                return

            # Восстанавливаем бэкап
            db_settings = settings.DATABASES['default']
            command = (
                f"pg_restore -h {db_settings['HOST']} -p {db_settings['PORT']} "
                f"-U {db_settings['USER']} -d {db_settings['NAME']} -v -c {backup.file.path}"
            )

            # Запускаем команду
            env = os.environ.copy()
            env['PGPASSWORD'] = db_settings['PASSWORD']
            subprocess.run(command, shell=True, check=True, env=env)

            messagebox.showinfo("Успех", "База данных успешно восстановлена из бэкапа")

            # Закрываем модальное окно и обновляем данные
            self.backup_modal.destroy()
            self.refresh_data()
        except Exception as e:
            self.show_error("Ошибка восстановления", str(e))

    def delete_selected_backup(self):
        """Удаление выбранного бэкапа"""
        try:
            selected_item = self.backup_table.selection()
            if not selected_item:
                messagebox.showwarning("Предупреждение", "Выберите бэкап для удаления")
                return

            # Получаем ID выбранного бэкапа
            backup_id = self.backup_table.item(selected_item[0])['values'][0]
            backup = DatabaseBackup.objects.get(id=backup_id)

            # Подтверждение действия
            if not messagebox.askyesno(
                    "Подтверждение",
                    f"Вы уверены, что хотите удалить бэкап от {backup.created_at.strftime('%d.%m.%Y %H:%M')}?"
            ):
                return

            # Удаляем бэкап
            backup.file.delete()
            backup.delete()

            # Обновляем список бэкапов
            self.load_backups_list()

            messagebox.showinfo("Успех", "Бэкап успешно удален")
        except Exception as e:
            self.show_error("Ошибка удаления", str(e))

    def get_current_user_account(self):
        """Получение текущего пользователя (заглушка - нужно реализовать в вашей системе аутентификации)"""
        # В реальной системе здесь должен быть код для получения текущего пользователя
        return Account.objects.first()  # Заглушка - возвращает первого пользователя

    def setup_context_menu(self):
        """Настройка контекстного меню для таблицы"""
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Обновить", command=self.refresh_data)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="Экспорт в Word", command=self.create_docx_file)
        self.context_menu.add_command(label="Экспорт в Excel", command=self.export_to_excel)

        self.table.bind("<Button-3>", self.show_context_menu)

    def show_context_menu(self, event):
        """Показать контекстное меню"""
        try:
            self.context_menu.tk_popup(event.x_root, event.y_root)
        finally:
            self.context_menu.grab_release()

    def resize_table(self, event):
        """Изменение размера таблицы при изменении окна"""
        try:
            self.table_frame.update_idletasks()
            table_width = self.table_frame.winfo_width()
            self.table.column("#0", width=0, stretch=tk.NO)

            # Вычисляем общую ширину всех столбцов
            total_width = sum(int(self.table.column(col)['width']) for col in self.columns)

            # Если общая ширина нулевая, устанавливаем равномерное распределение
            if total_width == 0:
                equal_width = table_width // len(self.columns)
                for col in self.columns:
                    self.table.column(col, width=equal_width)
                return

            # Распределяем ширину столбцов пропорционально
            for col in self.columns:
                current_width = int(self.table.column(col)['width'])
                new_width = int((current_width / total_width) * table_width)
                self.table.column(col, width=new_width)
        except Exception as e:
            print(f"Ошибка при изменении размера таблицы: {e}")

    def update_table(self, data):
        """Обновление данных в таблице"""
        for row in self.table.get_children():
            self.table.delete(row)
        for student in data:
            self.table.insert("", "end", values=student)

    def filter_students(self, search_text):
        """Фильтрация студентов по поисковому запросу"""
        if search_text.strip() == "":
            self.update_table(self.students_data)
        else:
            filtered = [student for student in self.students_data if any(
                search_text.lower() in str(value).lower() for value in student[1:]
            )]
            self.update_table(filtered)

    def filter_by_group(self, group_name):
        """Фильтрация студентов по группе"""
        if group_name == "Все группы":
            self.update_table(self.students_data)
        else:
            filtered = [student for student in self.students_data if student[6] == group_name]
            self.update_table(filtered)

    def refresh_data(self):
        """Обновление данных"""
        self.students_data = self.fetch_students()
        self.update_table(self.students_data)
        messagebox.showinfo("Обновлено", "Данные успешно обновлены")

    def open_add_student_modal(self):
        """Открытие модального окна для добавления студента"""
        self.modal = tk.Toplevel(self.root)
        self.modal.title("Добавить студента")
        self.modal.geometry("500x600")
        self.modal.resizable(False, False)
        self.modal.grab_set()

        # Форма добавления
        form_frame = ttk.Frame(self.modal)
        form_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)

        # Поля формы
        fields = [
            ("Фамилия*", "last_name"),
            ("Имя*", "first_name"),
            ("Отчество", "middle_name"),
            ("Телефон", "phone_number"),
            ("Метро", "metro_station"),
            ("Группа*", "group"),
            ("Организация", "organization"),
            ("Руководитель", "college_supervisor")
        ]

        self.entries = {}

        for i, (label_text, field_name) in enumerate(fields):
            frame = ttk.Frame(form_frame)
            frame.pack(fill=tk.X, pady=5)

            label = ttk.Label(frame, text=label_text, width=15, anchor=tk.W)
            label.pack(side=tk.LEFT, padx=5)

            if field_name == "group":
                groups = list(Group.objects.values_list('name', flat=True))
                entry = ttk.Combobox(frame, values=groups, state="readonly")
            elif field_name == "organization":
                orgs = list(Organization.objects.values_list('full_name', flat=True))
                entry = ttk.Combobox(frame, values=orgs)
            elif field_name == "college_supervisor":
                supervisors = list(CollegeSupervisor.objects.values_list('last_name', flat=True))
                entry = ttk.Combobox(frame, values=supervisors)
            else:
                entry = ttk.Entry(frame)

            entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
            self.entries[field_name] = entry

        # Кнопки
        button_frame = ttk.Frame(self.modal)
        button_frame.pack(pady=20, padx=20, fill=tk.X)

        cancel_btn = ttk.Button(button_frame, text="Отмена", command=self.modal.destroy)
        cancel_btn.pack(side=tk.RIGHT, padx=5)

        save_btn = ttk.Button(button_frame, text="Сохранить", command=self.save_student)
        save_btn.pack(side=tk.RIGHT, padx=5)

    def save_student(self):
        """Сохранение нового студента"""
        try:
            # Проверка обязательных полей
            if not self.entries['last_name'].get() or not self.entries['first_name'].get() or not self.entries[
                'group'].get():
                raise ValueError("Заполните обязательные поля (Фамилия, Имя, Группа)")

            data = {
                'last_name': self.entries['last_name'].get(),
                'first_name': self.entries['first_name'].get(),
                'middle_name': self.entries['middle_name'].get() or None,
                'phone_number': self.entries['phone_number'].get() or None,
                'metro_station': self.entries['metro_station'].get() or None,
            }

            # Обработка группы
            group_name = self.entries['group'].get()
            group = Group.objects.get(name=group_name)
            data['group'] = group

            # Обработка организации
            org_name = self.entries['organization'].get()
            if org_name:
                org, _ = Organization.objects.get_or_create(full_name=str(org_name).strip())
                data['organization'] = org

            # Обработка руководителя
            supervisor_name = self.entries['college_supervisor'].get()
            if supervisor_name:
                supervisor = CollegeSupervisor.objects.get(last_name=supervisor_name)
                data['college_supervisor'] = supervisor

            # Создание студента
            Intern.objects.create(**data)

            messagebox.showinfo("Успех", "Студент успешно добавлен")
            self.modal.destroy()
            self.refresh_data()
        except Exception as e:
            self.show_error("Ошибка при сохранении", str(e))

    def open_generate_file_modal(self):
        """Открытие модального окна для генерации файла"""
        self.gen_modal = tk.Toplevel(self.root)
        self.gen_modal.title("Создать отчет")
        self.gen_modal.geometry("600x500")
        self.gen_modal.resizable(False, False)
        self.gen_modal.grab_set()

        # Основной контент
        content_frame = ttk.Frame(self.gen_modal)
        content_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)

        # Выбор специальности
        specialty_frame = ttk.Frame(content_frame)
        specialty_frame.pack(fill=tk.X, pady=5)

        specialty_label = ttk.Label(specialty_frame, text="Специальность:", width=15, anchor=tk.W)
        specialty_label.pack(side=tk.LEFT, padx=5)

        self.specialties = list(Specialty.objects.values_list('name', flat=True))
        self.specialty_var = tk.StringVar()
        specialty_dropdown = ttk.Combobox(specialty_frame, textvariable=self.specialty_var,
                                          values=self.specialties, state="readonly")
        specialty_dropdown.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        specialty_dropdown.bind("<<ComboboxSelected>>", lambda e: self.update_practices_list())

        # Выбор практики
        practice_frame = ttk.Frame(content_frame)
        practice_frame.pack(fill=tk.X, pady=5)

        practice_label = ttk.Label(practice_frame, text="Практика:", width=15, anchor=tk.W)
        practice_label.pack(side=tk.LEFT, padx=5)

        self.practice_var = tk.StringVar()
        self.practice_dropdown = ttk.Combobox(practice_frame, textvariable=self.practice_var,
                                              state="readonly")
        self.practice_dropdown.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.practice_dropdown.bind("<<ComboboxSelected>>", lambda e: self.update_groups_list())

        # Фрейм для групп (теперь только для отображения)
        groups_display_frame = ttk.Frame(content_frame)
        groups_display_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        groups_label = ttk.Label(groups_display_frame, text="Группы для выбранной практики:")
        groups_label.pack(anchor=tk.W)

        self.groups_text = tk.Text(groups_display_frame, height=5, state=tk.DISABLED)
        self.groups_text.pack(fill=tk.BOTH, expand=True)

        # Кнопки
        button_frame = ttk.Frame(self.gen_modal)
        button_frame.pack(pady=20, padx=20, fill=tk.X)

        cancel_btn = ttk.Button(button_frame, text="Отмена", command=self.gen_modal.destroy)
        cancel_btn.pack(side=tk.RIGHT, padx=5)

        generate_btn = ttk.Button(button_frame, text="Создать", command=self.generate_practice_report)
        generate_btn.pack(side=tk.RIGHT, padx=5)

    def update_practices_list(self):
        """Обновление списка практик при выборе специальности"""
        try:
            specialty_name = self.specialty_var.get()
            if not specialty_name:
                return

            # Очищаем выбор практики и групп
            self.practice_var.set('')
            self.practice_dropdown['values'] = []
            self.groups_text.config(state=tk.NORMAL)
            self.groups_text.delete(1.0, tk.END)
            self.groups_text.config(state=tk.DISABLED)

            # Получаем специальность
            specialty = Specialty.objects.get(name=specialty_name)

            # Получаем все практики для этой специальности (через группы)
            practices = Practice.objects.filter(groups__specialty=specialty).distinct()

            # Формируем список названий практик
            practice_names = []
            for practice in practices:
                name = ""
                if practice.pp:
                    name = practice.pp
                elif practice.pm:
                    name = practice.pm
                elif practice.preddiplom:
                    name = "Преддипломная практика"

                if name:
                    practice_names.append(name)

            # Обновляем выпадающий список практик
            self.practice_dropdown['values'] = practice_names

        except Exception as e:
            self.show_error("Ошибка", f"Не удалось загрузить практики: {str(e)}")

    def update_groups_list(self):
        """Обновление списка групп при выборе практики"""
        try:
            specialty_name = self.specialty_var.get()
            practice_name = self.practice_var.get()
            if not specialty_name or not practice_name:
                return

            # Получаем специальность
            specialty = Specialty.objects.get(name=specialty_name)

            # Находим практику по названию
            practice = None
            if "Преддипломная" in practice_name:
                practice = Practice.objects.filter(
                    groups__specialty=specialty,
                    preddiplom=True
                ).first()
            else:
                practice = Practice.objects.filter(
                    groups__specialty=specialty,
                    pp=practice_name
                ).first()

                if not practice:
                    practice = Practice.objects.filter(
                        groups__specialty=specialty,
                        pm=practice_name
                    ).first()

            if not practice:
                raise ValueError("Практика не найдена")

            # Получаем группы для этой практики
            groups = practice.groups.all()

            # Обновляем текстовое поле с группами
            self.groups_text.config(state=tk.NORMAL)
            self.groups_text.delete(1.0, tk.END)
            self.groups_text.insert(tk.END, "\n".join([group.name for group in groups]))
            self.groups_text.config(state=tk.DISABLED)

        except Exception as e:
            self.show_error("Ошибка", f"Не удалось загрузить группы: {str(e)}")

    def generate_practice_report(self):
        """Генерация отчета для выбранной практики"""
        try:
            specialty_name = self.specialty_var.get()
            practice_name = self.practice_var.get()

            if not specialty_name or not practice_name:
                messagebox.showwarning("Предупреждение", "Выберите специальность и практику")
                return

            # Получаем специальность
            specialty = Specialty.objects.get(name=specialty_name)

            # Находим практику
            practice = None
            if "Преддипломная" in practice_name:
                practice = Practice.objects.filter(
                    groups__specialty=specialty,
                    preddiplom=True
                ).first()
            else:
                practice = Practice.objects.filter(
                    groups__specialty=specialty,
                    pp=practice_name
                ).first()

                if not practice:
                    practice = Practice.objects.filter(
                        groups__specialty=specialty,
                        pm=practice_name
                    ).first()

            if not practice:
                raise ValueError("Практика не найдена")

            # Получаем группы для этой практики
            groups = practice.groups.all()
            group_names = [group.name for group in groups]

            if not group_names:
                messagebox.showwarning("Предупреждение", "Нет групп для выбранной практики")
                return

            # Предлагаем выбрать место сохранения
            default_filename = f"Базы_практики_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Files", "*.docx")],
                initialfile=default_filename,
                title="Сохранить отчет как"
            )

            if not output_path:
                return

            # Генерируем документ
            self.create_bases_practice_doc(group_names, output_path, practice)

            messagebox.showinfo("Успех", f"Отчет успешно создан:\n{output_path}")
            self.gen_modal.destroy()
        except Exception as e:
            self.show_error("Ошибка при создании отчета", str(e))

    def create_bases_practice_doc(self, selected_groups, output_path, practice):
        """
        Создаёт документ Word с базой практик для выбранных групп
        с альбомной ориентацией, рамками таблицы и объединением ячеек организаций.
        """
        try:
            # Получаем данные студентов
            interns = Intern.objects.filter(group__name__in=selected_groups) \
                .select_related('group', 'organization', 'college_supervisor', 'group__specialty') \
                .prefetch_related('organization__organizationsupervisor_set') \
                .all()

            if not interns:
                raise ValueError("Нет студентов в выбранных группах")

            # Разделяем студентов: с организацией и без
            interns_with_org = [i for i in interns if i.organization]
            interns_without_org = [i for i in interns if not i.organization]

            # Сортировка студентов с организациями
            sorted_interns_with_org = sorted(
                interns_with_org,
                key=lambda x: (
                    'техникум' not in x.organization.full_name.lower(),
                    x.organization.full_name,
                    x.last_name
                )
            )

            # Общий список: сначала с организациями, потом без
            sorted_interns = sorted_interns_with_org + interns_without_org

            # Создаём документ
            doc = Document()

            # Устанавливаем альбомную ориентацию
            section = doc.sections[0]
            section.orientation = WD_ORIENTATION.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height

            # Добавляем "Приложение 1" (справа, курсив, 10pt)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run("Приложение 1")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.italic = True

            # Основной заголовок (жирный, 12pt)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Базы производственной практики (по профилю специальности)")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True

            # Информация о практике
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Добавляем название практики
            if practice.pp:
                run = p.add_run(f"{practice.pp} {practice.pm} ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            elif practice.preddiplom:
                run = p.add_run("Преддипломная практика ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

            # Добавляем специальность с квалификацией
            group = Group.objects.filter(name__in=selected_groups).first()
            if group and group.specialty:
                specialty = group.specialty
                run = p.add_run(f"Специальности {specialty.code} «{specialty.name}», классификация: "
                                f"{specialty.organization.name}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True

            # Период проведения
            p = doc.add_paragraph()
            run = p.add_run(
                f"Групп: {', '.join(selected_groups)} период проведения: {practice.schedule}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

            # Создаём таблицу с рамками
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'  # Добавляем рамки ко всем ячейкам

            # Настройка ширины столбцов (примерные значения)
            col_widths = [0.5, 3.0, 2.5, 1.0, 2.0, 2.5]  # В дюймах
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = Inches(width)

            # Заголовки таблицы (11pt, серый фон, выравнивание по центру)
            headers = [
                "№ п/п",
                "Наименование организации",
                "Ф.И.О. студента",
                "Группа",
                "Руководитель практической подготовки от техникума",
                "Руководитель практической подготовки от организации"
            ]

            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.bold = True  # Сделаем заголовки жирными

                # Серый фон для заголовков
                tc = hdr_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9D9D9')
                tcPr.append(shading)

            # Заполняем таблицу данными
            current_org = None
            org_start_row = 1  # Начинаем с 1, так как 0 - это заголовок
            org_rows = []

            current_org_supervisor = None
            org_supervisor_start_row = 1
            org_supervisor_rows = []

            for idx, intern in enumerate(sorted_interns, start=1):
                row_cells = table.add_row().cells

                # Устанавливаем шрифт Times New Roman 11pt для всех ячеек
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.0  # Межстрочный интервал 1.0
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                # № п/п
                row_cells[0].text = str(idx)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Организация (оставляем пустым для студентов без организации)
                org_name = intern.organization.full_name if intern.organization else ""
                row_cells[1].text = org_name
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # ФИО студента
                student_name = f"{intern.last_name} {intern.first_name} {intern.middle_name or ''}"
                row_cells[2].text = student_name
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Группа
                row_cells[3].text = intern.group.name
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Руководитель от техникума (с инициалами, без объединения)
                supervisor = ""
                if intern.college_supervisor:
                    supervisor = f"{intern.college_supervisor.last_name} {intern.college_supervisor.first_name[0]}.{intern.college_supervisor.middle_name[0] + '.' if intern.college_supervisor.middle_name else ''}"
                row_cells[4].text = supervisor
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Руководитель от организации (ФИО + телефон)
                org_supervisor_info = ""
                if intern.organization:
                    org_supervisors = intern.organization.organizationsupervisor_set.all()
                    if org_supervisors.exists():
                        supervisor = org_supervisors.first()
                        org_supervisor_info = f"{supervisor.last_name} {supervisor.first_name} {supervisor.middle_name or ''}"
                        if supervisor.phone_number:
                            org_supervisor_info += f"\nтел.: {supervisor.phone_number}"

                row_cells[5].text = org_supervisor_info
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[5].paragraphs[0].paragraph_format.wrap_text = True

                # Объединяем ячейки только для организаций и их руководителей
                if intern.organization:
                    # Проверяем, нужно ли объединять ячейки с организациями
                    if org_name != current_org:
                        if current_org is not None and len(org_rows) > 1:
                            # Объединяем ячейки для предыдущей организации
                            self.merge_cells(table, org_start_row, org_start_row + len(org_rows) - 1, 1)

                        current_org = org_name
                        org_start_row = idx
                        org_rows = [idx]
                    else:
                        org_rows.append(idx)

                    # Проверяем, нужно ли объединять ячейки с руководителями от организации
                    if org_supervisor_info != current_org_supervisor:
                        if current_org_supervisor is not None and len(org_supervisor_rows) > 1:
                            # Объединяем ячейки для предыдущего руководителя от организации
                            self.merge_cells(table, org_supervisor_start_row,
                                             org_supervisor_start_row + len(org_supervisor_rows) - 1, 5)

                        current_org_supervisor = org_supervisor_info
                        org_supervisor_start_row = idx
                        org_supervisor_rows = [idx]
                    else:
                        org_supervisor_rows.append(idx)
                else:
                    # Для студентов без организации не объединяем ячейки
                    pass

            # Объединяем ячейки для последней организации и руководителя от организации (если есть)
            if current_org is not None and len(org_rows) > 1:
                self.merge_cells(table, org_start_row, org_start_row + len(org_rows) - 1, 1)

            if current_org_supervisor is not None and len(org_supervisor_rows) > 1:
                self.merge_cells(table, org_supervisor_start_row,
                                 org_supervisor_start_row + len(org_supervisor_rows) - 1, 5)

            # Сохраняем документ
            doc.save(output_path)

        except Exception as e:
            raise Exception(f"Ошибка при создании документа: {str(e)}")

    def add_page_number(self, footer):
        """Добавляет номер страницы в нижний колонтитул"""
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

    def set_cell_shading(self, cell, fill_color):
        """Устанавливает цвет заливки ячейки"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_alignment(self, cell, align='left', vertical='center'):
        """
        Устанавливает выравнивание текста в ячейке
        :param align: 'left', 'center', 'right', 'both', 'distribute'
        :param vertical: 'top', 'center', 'bottom'
        """
        # Горизонтальное выравнивание
        cell.paragraphs[0].alignment = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'distribute': WD_ALIGN_PARAGRAPH.DISTRIBUTE
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

        # Вертикальное выравнивание
        tc_pr = cell._tc.get_or_add_tcPr()
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), vertical)
        tc_pr.append(v_align)

    def merge_cells(self, table, start_row, end_row, col):
        """
        Объединяет ячейки в таблице по вертикали
        :param table: таблица docx
        :param start_row: начальная строка (0-based)
        :param end_row: конечная строка (0-based)
        :param col: индекс столбца (0-based)
        """
        if start_row >= end_row:
            return

        cell = table.cell(start_row, col)
        cell_text = cell.text

        for row in range(start_row + 1, end_row + 1):
            next_cell = table.cell(row, col)
            cell.merge(next_cell)

        cell.text = cell_text
        self.set_cell_alignment(cell, align='left', vertical='center')

    def create_docx_file(self):
        """Создание Word-файла для текущего отображения в таблице"""
        try:
            # Получаем текущие данные из таблицы
            selected_items = self.table.selection()
            if selected_items:
                data = [self.table.item(item)['values'] for item in selected_items]
            else:
                data = [self.table.item(item)['values'] for item in self.table.get_children()]

            if not data:
                messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
                return

            # Предлагаем выбрать место сохранения
            default_filename = f"Студенты_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            output_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Files", "*.docx")],
                initialfile=default_filename,
                title="Сохранить документ как"
            )

            if not output_path:
                return

            # Создаем документ
            doc = Document()

            # Добавляем заголовок
            title = doc.add_paragraph("Список студентов")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Pt(16)
            title.runs[0].bold = True

            # Добавляем дату
            date_para = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Добавляем таблицу
            table = doc.add_table(rows=1, cols=len(self.columns))
            table.style = 'Table Grid'

            # Заголовки таблицы
            for i, col in enumerate(self.columns):
                cell = table.cell(0, i)
                cell.text = col
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Заполняем таблицу данными
            for row in data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Сохраняем документ
            doc.save(output_path)

            messagebox.showinfo("Успех", f"Документ успешно сохранен:\n{output_path}")

        except Exception as e:
            self.show_error("Ошибка при создании документа", str(e))

    def export_to_excel(self):
        """Экспорт данных в Excel"""
        try:
            # Получаем текущие данные из таблицы
            selected_items = self.table.selection()
            if selected_items:
                data = [self.table.item(item)['values'] for item in selected_items]
            else:
                data = [self.table.item(item)['values'] for item in self.table.get_children()]

            if not data:
                messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
                return

            # Создаем DataFrame
            df = pd.DataFrame(data, columns=self.columns)

            # Предлагаем выбрать место сохранения
            default_filename = f"Студенты_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
            output_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel Files", "*.xlsx")],
                initialfile=default_filename,
                title="Сохранить как Excel"
            )

            if not output_path:
                return

            # Сохраняем в Excel
            df.to_excel(output_path, index=False)

            messagebox.showinfo("Успех", f"Данные успешно экспортированы в:\n{output_path}")

        except Exception as e:
            self.show_error("Ошибка при экспорте в Excel", str(e))

    def upload_excel_file(self):
        """Загрузка данных из Excel файла"""
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Excel Files", "*.xlsx *.xls")],
                title="Выберите файл Excel для загрузки"
            )

            if not file_path:
                return

            # Чтение файла Excel
            df = pd.read_excel(file_path, sheet_name=None)

            for sheet_name, sheet_data in df.items():
                # Извлекаем название группы
                group_name = sheet_name.replace("Группа ", "").strip()

                # Проверяем существование группы
                existing_group = Group.objects.filter(name=group_name).first()
                if existing_group:
                    continue  # Пропускаем существующие группы

                # Определяем специальность
                specialty_code = "09.02.07" if group_name.startswith(("П50", "П")) else None
                specialty_code = "09.02.01" if group_name.startswith(("Э50", "Э")) else None
                if not specialty_code:
                    continue

                specialty = Specialty.objects.filter(code=specialty_code).first()
                if not specialty:
                    continue

                # Создаем группу
                group = Group.objects.create(name=group_name, specialty=specialty)

                # Обрабатываем данные студентов
                for _, row in sheet_data.iterrows():
                    # Проверяем обязательные поля
                    if pd.isna(row.get("Фамилия")) or pd.isna(row.get("Имя")):
                        continue

                    # Создаем студента
                    intern_data = {
                        'last_name': row.get("Фамилия", ""),
                        'first_name': row.get("Имя", ""),
                        'middle_name': row.get("Отчество", None),
                        'phone_number': row.get("Телефон", None),
                        'metro_station': row.get("Метро", None),
                        'group': group,
                    }

                    # Обработка организации
                    org_name = row.get("Организация", None)
                    if org_name and not pd.isna(org_name):
                        org, _ = Organization.objects.get_or_create(full_name=str(org_name).strip())
                        intern_data['organization'] = org

                    Intern.objects.create(**intern_data)

            messagebox.showinfo("Успех", "Данные успешно загружены из Excel")
            self.refresh_data()

        except Exception as e:
            self.show_error("Ошибка при загрузке Excel", str(e))

    def show_error(self, title, message):
        """Показать сообщение об ошибке"""
        messagebox.showerror(title, f"{message}\n\nПодробности:\n{traceback.format_exc()}")


if __name__ == "__main__":
    root = tk.Tk()
    app = StudentManagementApp(root)
    root.mainloop()