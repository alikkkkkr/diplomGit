from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_bases_practice_doc(output_path):
    # Создаем новый документ
    doc = Document()

    # Добавляем заголовок "Приложение 1"
    appendix = doc.add_paragraph()
    appendix.add_run("Приложение 1").bold = True
    appendix.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем заголовок "Базы производственной практики (по профилю специальности)"
    title = doc.add_paragraph()
    title.add_run("Базы производственной практики (по профилю специальности)").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем текст с информацией о специальности и периоде проведения
    info = doc.add_paragraph()
    info.add_run("ПП.04.01 «Сопровождение и обслуживание программного обеспечения компьютерных систем» профессионального модуля ПМ.04 «Сопровождение и обслуживание программного обеспечения компьютерных систем» ")
    info.add_run("Специальности 09.02.07 «Информационные системы и программирование», классификация: Программист").bold = True
    info.add_run("\nГрупп: П50-1-21, П50-2-21, П50-3-21, П50-4-21, П50-5-21, П50-6-21, П50-7-21, П50-8-21, П50-9-21 период проведения: c 10.02.2025 по 12.04.2025")

    # Добавляем таблицу
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    # Настройка ширины столбцов
    for i, width in enumerate([0.5, 3, 3, 1.5, 2, 3]):
        col = table.columns[i]
        col.width = Pt(width * 72)  # Преобразуем дюймы в пункты (1 дюйм = 72 пункта)

    # Заголовки таблицы
    headers = ["№ п/п", "Наименование организации", "Ф.И.О. студента", "Группа", "Руководитель практической подготовки от техникума", "Руководитель практической подготовки от организации"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Пример данных для таблицы (замените на свои данные)
    data = [
        ["1", "ООО \"02 КЛАУД\"", "Астанин Илья Игоревич", "П50-1-21", "Соколова Л.А.", "Будкин Александр Владимирович"],
        ["2", "АО \"Уральский завод гражданской авиации\"", "Брызгалов Михаил Игоревич", "П50-1-21", "Соколова Л.А.", "Кругликов Андрей Игоревич 8-925-727-03-41"],
        # Добавьте остальные строки данных
    ]

    # Заполнение таблицы данными
    for row in data:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = cell_value
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Сохраняем документ
    doc.save(output_path)

# Пример вызова функции
create_bases_practice_doc("Базы_практик_шаблон.docx")
